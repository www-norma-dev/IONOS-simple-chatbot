import os
import tempfile
import logging
from typing import Iterable, List, Optional

import requests
from fastapi import HTTPException

from pdfminer.high_level import extract_text as pdf_extract_text  # type: ignore
from docx import Document  # type: ignore


logger = logging.getLogger("chatbot-server")


class DocumentCollector:
    """
    Collect text content from local files/directories and from direct file URLs.

    Supported extensions: .pdf, .docx, .txt
    """

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}

    def __init__(self, chunk_size: int = 500, max_chunk_count: int = 256):
        self.chunk_size = chunk_size
        self.max_chunk_count = max_chunk_count

    async def collect_documents(self, sources: Iterable[str]) -> List[str]:
        """
        Collect and return text chunks from the given sources.

        A source can be a local file path, a directory path (will be scanned recursively),
        or a direct URL to a file (http/https).

        Args:
            sources: Iterable of paths or URLs

        Returns:
            List of text chunks

        Raises:
            HTTPException: If none of the sources can be processed
        """
        if not sources:
            return []

        aggregated_chunks: List[str] = []
        any_success = False

        for source in sources:
            source = source.strip()
            try:
                if self._is_url(source):
                    logger.info("Downloading document from URL: %s", source)
                    text = self._extract_text_from_url(source)
                else:
                    text = self._extract_text_from_path(source)

                if not text:
                    logger.warning("No text extracted from source: %s", source)
                    continue

                any_success = True
                aggregated_chunks.extend(self._chunk_text(text))

                if len(aggregated_chunks) >= self.max_chunk_count:
                    logger.info(
                        "Reached max chunk count (%d); truncating results",
                        self.max_chunk_count,
                    )
                    aggregated_chunks = aggregated_chunks[: self.max_chunk_count]
                    break
            except HTTPException:
                raise
            except Exception as exc:
                logger.error("Failed to process source %s: %s", source, exc)

        if not any_success and not aggregated_chunks:
            raise HTTPException(status_code=400, detail="No valid documents found in sources")

        logger.info("Collected %d chunks from %d source(s)", len(aggregated_chunks), len(list(sources)))
        return aggregated_chunks

    # -----------------------------
    # Helpers
    # -----------------------------
    @staticmethod
    def _is_url(value: str) -> bool:
        return value.lower().startswith(("http://", "https://"))

    def _extract_text_from_path(self, path: str) -> str:
        if not os.path.exists(path):
            raise HTTPException(status_code=404, detail=f"Path not found: {path}")

        if os.path.isdir(path):
            texts: List[str] = []
            for root, _dirs, files in os.walk(path):
                for filename in files:
                    if self._is_supported(filename):
                        file_path = os.path.join(root, filename)
                        file_text = self._extract_text_from_file(file_path)
                        if file_text:
                            texts.append(file_text)
            return "\n\n".join(texts)
        else:
            if not self._is_supported(path):
                raise HTTPException(
                    status_code=400,
                    detail=(
                        f"Unsupported file type for path: {path}. Supported: {sorted(self.SUPPORTED_EXTENSIONS)}"
                    ),
                )
            return self._extract_text_from_file(path)

    def _extract_text_from_url(self, url: str) -> str:
        # Stream download into a temporary file, then parse.
        response = requests.get(url, timeout=60)
        try:
            response.raise_for_status()
        except Exception as exc:  # requests.HTTPError
            raise HTTPException(status_code=500, detail=f"Failed to download {url}: {exc}")

        filename = self._infer_filename_from_url(url, response.headers.get("content-type"))
        if not self._is_supported(filename):
            raise HTTPException(
                status_code=400,
                detail=(
                    f"Unsupported file type from URL: {filename}. Supported: {sorted(self.SUPPORTED_EXTENSIONS)}"
                ),
            )

        suffix = os.path.splitext(filename)[1]
        with tempfile.NamedTemporaryFile(delete=True, suffix=suffix) as tmp:
            tmp.write(response.content)
            tmp.flush()
            return self._extract_text_from_file(tmp.name)

    def _extract_text_from_file(self, file_path: str) -> str:
        _, ext = os.path.splitext(file_path.lower())
        try:
            if ext == ".pdf":
                return self._extract_text_from_pdf(file_path)
            if ext == ".docx":
                return self._extract_text_from_docx(file_path)
            if ext == ".txt":
                return self._extract_text_from_txt(file_path)
        except HTTPException:
            raise
        except Exception as exc:
            raise HTTPException(status_code=500, detail=f"Failed to parse {file_path}: {exc}")

        raise HTTPException(status_code=400, detail=f"Unsupported file type: {ext}")

    @staticmethod
    def _extract_text_from_pdf(file_path: str) -> str:
        text = pdf_extract_text(file_path) or ""
        return text.strip()

    @staticmethod
    def _extract_text_from_docx(file_path: str) -> str:
        doc = Document(file_path)
        parts: List[str] = []
        for paragraph in doc.paragraphs:
            if paragraph.text and paragraph.text.strip():
                parts.append(paragraph.text.strip())
        # Include table text
        for table in doc.tables:
            for row in table.rows:
                cells_text = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                if cells_text:
                    parts.append("\t".join(cells_text))
        return "\n".join(parts).strip()

    @staticmethod
    def _extract_text_from_txt(file_path: str) -> str:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read().strip()

    def _chunk_text(self, text: str) -> List[str]:
        if not text:
            return []

        chunks: List[str] = []
        for i in range(0, len(text), self.chunk_size):
            chunk = text[i : i + self.chunk_size]
            if chunk.strip():
                chunks.append(chunk)
            if len(chunks) >= self.max_chunk_count:
                break
        return chunks

    def _is_supported(self, path_or_name: str) -> bool:
        _, ext = os.path.splitext(path_or_name.lower())
        return ext in self.SUPPORTED_EXTENSIONS

    @staticmethod
    def _infer_filename_from_url(url: str, content_type: Optional[str]) -> str:
        # Try to grab filename from URL path
        name = os.path.basename(url.split("?")[0].split("#")[0])
        if name and "." in name:
            return name
        # Fallback to content type
        if content_type:
            mapping = {
                "application/pdf": "download.pdf",
                "application/msword": "download.doc",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "download.docx",
                "text/plain": "download.txt",
            }
            return mapping.get(content_type.split(";")[0].strip(), "download.bin")
        return "download.bin"


